from flask import Flask, request, jsonify, render_template
import os
import requests
import time
import smtplib
from email.mime.text import MIMEText
from docx import Document
from email.mime.base import MIMEBase
from email import encoders
from email.mime.multipart import MIMEMultipart

ASSEMBLY_API_KEY = '1cfced456ed0499c8d67290d2a1272e1'  # Your actual AssemblyAI key 
OPENROUTER_API_KEY = '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'  # Your actual OpenRouter API key 
SENDER_EMAIL = 'serdar1392@gmail.com'
SENDER_PASSWORD = 'uvxr mzqo jbue nmdi'

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

def transcribe_with_assemblyai(file_path):
    headers = {'authorization': ASSEMBLY_API_KEY}
    with open(file_path, 'rb') as f:
        upload_response = requests.post('https://api.assemblyai.com/v2/upload', headers=headers, files={'file': f})
    upload_url = upload_response.json()['upload_url']

    transcript_request = {'audio_url': upload_url, 'auto_chapters': False}
    transcript_response = requests.post('https://api.assemblyai.com/v2/transcript', json=transcript_request, headers=headers)
    transcript_id = transcript_response.json()['id']

    while True:
        polling_response = requests.get(f'https://api.assemblyai.com/v2/transcript/{transcript_id}', headers=headers)
        status = polling_response.json()['status']
        if status == 'completed':
            return polling_response.json()['text']
        elif status == 'error':
            return f"Error: {polling_response.json()['error']}"
        time.sleep(3)

def summarize_text_with_openrouter(text):
    headers = {
    "Authorization": f"Bearer {os.environ['OPENROUTER_API_KEY']}",
    "Content-Type": "application/json"
    }


    data = {
        "model": "anthropic/claude-3-sonnet",
        "max_tokens": 1000,
        "messages": [
            {"role": "user", "content": f"Summarize this: {text}"}
        ]
    }

    try:
        response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=data)
        print("\U0001F50D OpenRouter Raw Response:", response.text)
        result = response.json()
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        print("❌ Summarization Error:", e)
        return "Summarization failed."

def save_transcription_to_docx(text, filename):
    doc = Document()
    doc.add_heading('Transcription', level=1)
    doc.add_paragraph(text)
    doc.save(filename)
    return filename

def send_email(subject, recipient, attachment_paths=None):
    msg = MIMEMultipart()
    msg['Subject'] = subject
    msg['From'] = SENDER_EMAIL
    msg['To'] = recipient

    msg.attach(MIMEText("Please find the transcription and summary attached.", 'plain'))

    if attachment_paths:
        for path in attachment_paths:
            with open(path, 'rb') as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(path)}')
                msg.attach(part)

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'audio' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['audio']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    transcription = transcribe_with_assemblyai(file_path)
    summary = summarize_text_with_openrouter(transcription)

    docx_transcription = save_transcription_to_docx(transcription, 'transcription_full.docx')
    docx_summary = save_transcription_to_docx(summary, 'transcription_summary.docx')

    recipients = request.form.get('recipient', '').split(',')
    for email in recipients:
        email = email.strip()
        if email:
            send_email(
                subject='Transcription Summary & Full Text',
                recipient=email,
                attachment_paths=[docx_summary, docx_transcription]
            )

    return jsonify({'message': 'Transcription complete', 'text': transcription})

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))  # Railway injects PORT
    app.run(host='0.0.0.0', port=port)
